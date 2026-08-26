from datetime import date
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from app.document_reader import DocumentReadError, _extract_pdf_page_text, read_document
from app.domain import CandidateProfile, DocumentType, ProcessingStage, SessionStatus
from app.profile_extractor import ProfileExtractor
from app.session_store import SessionStore
from app.worker import DocumentWorker


class UnusedProfileGenerator:
    async def generate(self, **kwargs):
        raise AssertionError("Profile extraction should not run in this test")


class LayoutBlankPage:
    def extract_text(self, *args, **kwargs) -> str:
        if kwargs.get("extraction_mode") == "layout":
            return ""
        return "Readable CV text recovered by normal PDF extraction mode."


def write_pdf(path: Path, *lines: str) -> None:
    document = canvas.Canvas(str(path))
    y = 760
    for line in lines:
        document.drawString(72, y, line)
        y -= 20
    document.save()


def test_reads_text_pdf_with_page_evidence(tmp_path) -> None:
    source = tmp_path / "resume.pdf"
    write_pdf(
        source,
        "Alex Morgan - Platform Engineer",
        "Built reliable Python and Kubernetes services for six years.",
    )

    extracted = read_document(
        source,
        resume_id=uuid4(),
        max_pages=25,
        max_docx_uncompressed_bytes=50 * 1024 * 1024,
    )

    assert extracted.document_type == DocumentType.PDF
    assert extracted.page_count == 1
    assert extracted.blocks[0].page_number == 1
    assert "Kubernetes" in extracted.combined_text()


def test_falls_back_when_pdf_layout_mode_discards_text() -> None:
    text = _extract_pdf_page_text(LayoutBlankPage())  # type: ignore[arg-type]

    assert "Readable CV text" in text


def test_reads_docx_paragraphs_and_tables(tmp_path) -> None:
    source = tmp_path / "resume.docx"
    document = Document()
    document.add_heading("Alex Morgan", level=1)
    document.add_paragraph("Senior data analyst with seven years of experience.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "SQL, Python, Power BI"
    document.save(source)

    extracted = read_document(
        source,
        resume_id=uuid4(),
        max_pages=25,
        max_docx_uncompressed_bytes=50 * 1024 * 1024,
    )

    assert extracted.document_type == DocumentType.DOCX
    assert extracted.page_count is None
    assert {block.kind for block in extracted.blocks} == {"paragraph", "table"}
    assert "Power BI" in extracted.combined_text()


def test_rejects_image_only_pdf_with_clear_message(tmp_path) -> None:
    source = tmp_path / "scanned.pdf"
    document = canvas.Canvas(str(source))
    document.rect(72, 700, 200, 50)
    document.save()

    with pytest.raises(DocumentReadError, match="Scanned CVs are not supported"):
        read_document(
            source,
            resume_id=uuid4(),
            max_pages=25,
            max_docx_uncompressed_bytes=50 * 1024 * 1024,
        )


@pytest.mark.asyncio
async def test_worker_writes_extraction_then_deletes_original(tmp_path) -> None:
    source = tmp_path / "source.pdf"
    write_pdf(
        source,
        "Alex Morgan - Backend Engineer",
        "Designed Python APIs and operated Docker services in production.",
    )
    store = SessionStore(tmp_path / "sessions", ttl_seconds=600)
    manifest = store.create(
        original_filename="resume.pdf",
        content_type="application/pdf",
        content=source.read_bytes(),
    )
    worker = DocumentWorker(
        store=store,
        max_document_pages=25,
        max_docx_uncompressed_bytes=50 * 1024 * 1024,
        poll_seconds=0.01,
        lease_seconds=30,
        profile_extractor=ProfileExtractor(UnusedProfileGenerator()),
    )

    await worker.process_session(manifest.match_session_id)

    completed = store.read_manifest(manifest.match_session_id)
    assert completed.status == SessionStatus.PROCESSING
    assert completed.stage == ProcessingStage.BUILDING_PROFILE
    assert completed.raw_upload_deleted_at is not None
    assert not store.upload_path(completed).exists()
    assert store.artifact_exists(completed.match_session_id, "extracted-document.json")

    # Simulate a crash after the atomic profile write but before the manifest advance.
    profile = CandidateProfile(
        profile_id=uuid4(),
        resume_id=completed.resume_id,
        extractor_version="crash-recovery-test",
        experience_as_of=date(2026, 8, 26),
        total_experience_months=0,
        featured_experience_id=None,
        experiences=[],
        skills=[],
        qualifications=[],
        education=[],
    )
    store.write_model_artifact(
        completed.match_session_id,
        "candidate-profile.json",
        profile,
    )

    await worker.process_session(completed.match_session_id)

    recovered = store.read_manifest(completed.match_session_id)
    assert recovered.stage == ProcessingStage.FINDING_MATCHES
    assert recovered.profile_filename == "candidate-profile.json"
